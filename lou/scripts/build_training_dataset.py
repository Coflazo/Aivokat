#!/usr/bin/env python3
"""Build auditable weak-label training data from the Siemens sample documents."""

from __future__ import annotations

import json
import math
import re
import sys
from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from backend.services.parser import enrich_rules_with_word_doc, parse_excel_playbook


PROJECT_ROOT = Path(__file__).resolve().parent.parent
SIEMENS_DOCS = PROJECT_ROOT.parent / "Siemens Sample Documents"
PLAYBOOK_XLSX = SIEMENS_DOCS / "Sample NDA Playbook.csv.xlsx"
PLAYBOOK_DOCX = SIEMENS_DOCS / "Sample NDA Playbook.docx"
TRAINING_DIR = PROJECT_ROOT / "backend" / "data" / "training"

CONTRACTS = [
    SIEMENS_DOCS / "Sample Standard NDA.docx",
    *sorted((SIEMENS_DOCS / "Sample NDAs" / "Standard NDAs negotiated").glob("*.docx")),
    *sorted((SIEMENS_DOCS / "Sample NDAs" / "Customer NDAs").glob("*.docx")),
]

SYNTHETIC_QUESTION_TEMPLATES = [
    "What is our position on {topic}?",
    "Can we accept changes to {topic}?",
    "When do we escalate {topic}?",
    "What is the fallback for {topic}?",
]

CONTRADICT_CUES = [
    "unlimited",
    "uncapped",
    "all information",
    "perpetual",
    "penalty",
    "exclusive jurisdiction",
    "prior written consent",
    "without exception",
    "shall be liable for all",
]


def main() -> None:
    TRAINING_DIR.mkdir(parents=True, exist_ok=True)
    rules = enrich_rules_with_word_doc(parse_excel_playbook(PLAYBOOK_XLSX), PLAYBOOK_DOCX)
    rule_texts = {rule["rule_id"]: rule_to_text(rule) for rule in rules}

    retriever_rows: list[dict] = []
    classifier_rows: list[dict] = []
    audit_rows: list[dict] = []

    for rule in rules:
        negatives = hard_negatives(rule, rules)
        for template in SYNTHETIC_QUESTION_TEMPLATES:
            query = template.format(topic=rule["topic"])
            retriever_rows.append({
                "query": query,
                "positive_rule_id": rule["rule_id"],
                "positive_text": rule_texts[rule["rule_id"]],
                "negative_rule_ids": negatives,
                "source": "synthetic_question",
            })

    for path in CONTRACTS:
        family = contract_family(path)
        for clause in extract_clause_candidates(path):
            if len(clause.split()) < 6:
                continue
            match = best_rule_match(clause, rules)
            if not match:
                continue
            rule, score = match
            label, confidence, rationale = weak_label(clause, rule, score, family)
            negative_ids = hard_negatives(rule, rules)

            retriever_rows.append({
                "query": clause,
                "positive_rule_id": rule["rule_id"],
                "positive_text": rule_texts[rule["rule_id"]],
                "negative_rule_ids": negative_ids,
                "source": path.name,
            })
            classifier_rows.append({
                "clause_text": clause,
                "rule_id": rule["rule_id"],
                "rule_text": rule_texts[rule["rule_id"]],
                "label": label,
                "source_document": path.name,
                "document_family": family,
                "lexical_similarity": score,
                "label_source": "weak_heuristic",
                "confidence": confidence,
                "rationale": rationale,
            })
            audit_rows.append({
                "source_document": path.name,
                "clause_text": clause,
                "matched_rule_id": rule["rule_id"],
                "matched_topic": rule["topic"],
                "label": label,
                "confidence": confidence,
                "rationale": rationale,
            })

    write_jsonl(TRAINING_DIR / "retriever_pairs.jsonl", retriever_rows)
    write_jsonl(TRAINING_DIR / "classifier_pairs.jsonl", classifier_rows)
    write_jsonl(TRAINING_DIR / "label_audit.jsonl", audit_rows)

    print(f"rules: {len(rules)}")
    print(f"retriever rows: {len(retriever_rows)}")
    print(f"classifier rows: {len(classifier_rows)}")
    print(f"classifier labels: {dict(Counter(row['label'] for row in classifier_rows))}")
    print(f"output: {TRAINING_DIR}")


def extract_clause_candidates(path: Path) -> list[str]:
    doc = Document(str(path))
    candidates: list[str] = []
    current_heading = ""
    for paragraph in doc.paragraphs:
        text = clean_space(paragraph.text)
        if not text:
            continue
        if is_heading(text):
            current_heading = text
            continue
        if looks_like_clause(text):
            candidates.append(clean_space(f"{current_heading}. {text}") if current_heading else text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_space(cell.text) for cell in row.cells if clean_space(cell.text)]
            if cells:
                candidates.append(" | ".join(cells))
    return dedupe(candidates)


def best_rule_match(clause: str, rules: list[dict]) -> tuple[dict, float] | None:
    clause_tokens = weighted_tokens(clause)
    if not clause_tokens:
        return None
    scored: list[tuple[float, dict]] = []
    for rule in rules:
        rule_tokens = weighted_tokens(rule_to_text(rule))
        score = cosine_counter(clause_tokens, rule_tokens)
        topic_tokens = set(tokenize(rule["topic"]))
        if topic_tokens & set(clause_tokens):
            score += 0.15
        scored.append((score, rule))
    score, rule = max(scored, key=lambda item: item[0])
    return (rule, min(score, 1.0)) if score >= 0.08 else None


def weak_label(clause: str, rule: dict, score: float, family: str) -> tuple[str, float, str]:
    lower = clause.lower()
    red_line = (rule.get("red_line") or "").lower()
    standard = (rule.get("standard_position") or "").lower()
    fallback = (rule.get("fallback_position") or "").lower()

    if red_line and lexical_overlap(lower, red_line) >= 0.22:
        return "CONTRADICTS", 0.86, "Clause overlaps with the rule red line."
    if any(cue in lower for cue in CONTRADICT_CUES) and family == "customer":
        return "CONTRADICTS", 0.72, "Customer-side clause contains adversarial red-line language."
    if standard and lexical_overlap(lower, standard) >= 0.28:
        return "CONFIRMS", 0.82, "Clause overlaps with the preferred position."
    if fallback and lexical_overlap(lower, fallback) >= 0.22:
        return "EXTENDS", 0.76, "Clause is closest to fallback or variant language."
    if score < 0.16:
        return "NEW_RULE", 0.68, "No strong playbook match found."
    if family == "standard":
        return "CONFIRMS", 0.70, "Standard NDA clause weakly maps to an existing rule."
    return "EXTENDS", 0.62, "Negotiated clause is related but not close enough to confirm the rule."


def hard_negatives(rule: dict, rules: list[dict], count: int = 3) -> list[str]:
    topic_tokens = set(tokenize(rule["topic"]))
    scored = []
    for candidate in rules:
        if candidate["rule_id"] == rule["rule_id"]:
            continue
        overlap = len(topic_tokens & set(tokenize(candidate["topic"])))
        scored.append((overlap, candidate["rule_id"]))
    return [rule_id for _, rule_id in sorted(scored)[:count]]


def rule_to_text(rule: dict) -> str:
    parts = [
        f"Topic: {rule['topic']}",
        f"Category: {rule.get('category', 'Other')}",
        f"Standard Position: {rule.get('standard_position', '')}",
    ]
    if rule.get("fallback_position"):
        parts.append(f"Fallback Position: {rule['fallback_position']}")
    if rule.get("red_line"):
        parts.append(f"Red Line: {rule['red_line']}")
    if rule.get("decision_logic"):
        parts.append(f"Escalation Trigger: {rule['decision_logic']}")
    if rule.get("reasoning"):
        parts.append(f"Why It Matters: {rule['reasoning']}")
    return "\n".join(parts)


def contract_family(path: Path) -> str:
    if path.name.startswith("Sample Standard"):
        return "standard"
    if path.name.startswith("A"):
        return "a_series"
    if path.name.startswith("B"):
        return "customer"
    return "unknown"


def looks_like_clause(text: str) -> bool:
    return bool(re.match(r"^(\d+(\.\d+)?\.?|[A-Z]\.|ARTICLE\s+\d+|[a-z]\))\s+", text)) or len(text.split()) >= 14


def is_heading(text: str) -> bool:
    if len(text.split()) > 8:
        return False
    return text.isupper() or bool(re.match(r"^(ARTICLE\s+\d+|[A-Z]\.)", text))


def weighted_tokens(text: str) -> Counter:
    return Counter(tokenize(text))


def tokenize(text: str) -> list[str]:
    stop = {
        "the", "and", "or", "of", "to", "a", "an", "in", "for", "with", "by",
        "party", "parties", "shall", "may", "this", "that", "it", "is", "are",
    }
    return [token for token in re.findall(r"[a-z0-9]+", text.lower()) if token not in stop and len(token) > 2]


def cosine_counter(left: Counter, right: Counter) -> float:
    common = set(left) & set(right)
    numerator = sum(left[token] * right[token] for token in common)
    left_norm = math.sqrt(sum(value * value for value in left.values()))
    right_norm = math.sqrt(sum(value * value for value in right.values()))
    if not left_norm or not right_norm:
        return 0.0
    return numerator / (left_norm * right_norm)


def lexical_overlap(left: str, right: str) -> float:
    left_tokens = set(tokenize(left))
    right_tokens = set(tokenize(right))
    if not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / len(right_tokens)


def clean_space(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def dedupe(items: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for item in items:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        output.append(item)
    return output


def write_jsonl(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=True) + "\n")


if __name__ == "__main__":
    main()
