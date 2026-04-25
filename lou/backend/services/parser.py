import re
from pathlib import Path

import openpyxl
from docx import Document

EXTRACTION_SYSTEM = """You are a legal knowledge extraction specialist.
Extract negotiation playbook rules into structured JSON.
Return only valid JSON."""

EXTRACTION_USER_TEMPLATE = """Extract all negotiation rules from this playbook content.

Return:
{{
  "rules": [
    {{
      "rule_id": "snake_case_unique_id",
      "topic": "Human-readable topic name",
      "category": "Financial Risk | Confidentiality | IP & Data | Governance | Payment | Termination | Warranties | Dispute Resolution | Other",
      "standard_position": "What we want as the default position",
      "fallback_position": "What we can accept under pressure, or null",
      "red_line": "What we never accept, or null",
      "reasoning": "Plain-English reason",
      "suggested_language": null,
      "decision_logic": "When to escalate, or null"
    }}
  ]
}}

Keep each negotiation topic as one rule object. Do not split standard,
fallback, and red-line positions into separate rules.

PLAYBOOK CONTENT:
{content}"""


def parse_excel_playbook(file_path: str | Path) -> list[dict]:
    """Parse the Siemens NDA playbook directly.

    The real file is named "Sample NDA Playbook.csv.xlsx" and has eight
    columns:
    0 Clause #, 1 Clause Name, 2 Why It Matters, 3 Preferred Position,
    4 Fallback 1, 5 Fallback 2, 6 Red Line, 7 Escalation Trigger.
    """
    wb = openpyxl.load_workbook(file_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    header = [str(cell or "").strip() for cell in rows[0]]
    expected = [
        "Clause #",
        "Clause Name",
        "Why It Matters (Summary)",
        "Preferred Position",
        "Fallback 1",
        "Fallback 2",
        "Red Line",
        "Escalation Trigger",
    ]
    if header[:8] != expected:
        raise ValueError(f"Unexpected playbook columns: {header[:8]}")

    rules: list[dict] = []
    for row in rows[1:]:
        if not row or not row[1]:
            continue

        clause_name = str(row[1]).strip()
        why = str(row[2]).strip() if row[2] else ""
        preferred = str(row[3]).strip() if row[3] else ""
        fallback_1 = str(row[4]).strip() if row[4] else ""
        fallback_2 = str(row[5]).strip() if row[5] else ""
        red_line = str(row[6]).strip() if row[6] else ""
        escalation = str(row[7]).strip() if row[7] else ""

        fallbacks = []
        if fallback_1:
            fallbacks.append(f"Fallback 1: {fallback_1}")
        if fallback_2:
            fallbacks.append(f"Fallback 2: {fallback_2}")

        rules.append({
            "rule_id": _to_snake_case(clause_name),
            "topic": clause_name,
            "category": _infer_category(clause_name),
            "rule_type": "standard",
            "standard_position": preferred,
            "fallback_position": "\n".join(fallbacks) if fallbacks else None,
            "red_line": red_line or None,
            "reasoning": why,
            "suggested_language": None,
            "decision_logic": escalation or None,
        })

    return rules


def enrich_rules_with_word_doc(rules: list[dict], word_path: str | Path) -> list[dict]:
    if not Path(word_path).exists():
        return rules

    text = _extract_docx_text(word_path)
    enriched: list[dict] = []
    for rule in rules:
        section = _find_section_text(text, rule["topic"])
        if section and section not in rule.get("reasoning", ""):
            rule = {**rule, "reasoning": f"{rule.get('reasoning', '')}\n\n{section}".strip()}
        enriched.append(rule)
    return enriched


async def parse_playbook(file_path: str, filename: str) -> list[dict]:
    suffix = Path(filename).suffix.lower()
    name = Path(filename).name
    if suffix in (".xlsx", ".xls") and name == "Sample NDA Playbook.csv.xlsx":
        return parse_excel_playbook(file_path)

    if suffix in (".docx", ".doc"):
        content = _extract_docx_text(file_path)
    elif suffix in (".xlsx", ".xls", ".csv"):
        content = _extract_xlsx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    all_rules: list[dict] = []
    seen_ids: set[str] = set()
    from backend.services.llm import complete_json

    for chunk in _chunks(content, 12000):
        result = await complete_json(
            EXTRACTION_SYSTEM,
            EXTRACTION_USER_TEMPLATE.format(content=chunk),
            max_tokens=4000,
        )
        rules = result.get("rules", result) if isinstance(result, dict) else result
        for rule in rules:
            rid = str(rule.get("rule_id", "")).strip()
            if not rid or rid in seen_ids:
                continue
            seen_ids.add(rid)
            all_rules.append(rule)
    return all_rules


async def extract_contract_clauses(file_path: str | Path) -> list[dict]:
    suffix = Path(file_path).suffix.lower()
    if suffix in (".docx", ".doc"):
        text = _extract_docx_text(file_path)
    elif suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        raise ValueError(f"Unsupported contract format: {suffix}")

    from backend.services.llm import complete_json

    result = await complete_json(
        """You extract individual negotiation clauses from contract documents.
Return only valid JSON.""",
        """Return a JSON object with a "clauses" array. Each item must have:
topic, clause_text, implied_position.

Focus on confidentiality, recipients, return/destruction, liability, penalties,
indemnification, IP, term, governing law, dispute resolution, language, and
signatures. Ignore definitions, recitals, and signature blocks.

CONTRACT TEXT:
""" + text[:14000],
        max_tokens=3000,
    )
    return result.get("clauses", []) if isinstance(result, dict) else []


def _extract_docx_text(path: str | Path) -> str:
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx_text(path: str | Path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                parts.append(" | ".join(str(cell or "") for cell in row))
    return "\n".join(parts)


def _find_section_text(full_text: str, topic: str) -> str:
    lines = full_text.splitlines()
    topic_key = re.sub(r"[^a-z0-9]+", " ", topic.lower()).strip()
    best_idx = -1
    for idx, line in enumerate(lines):
        line_key = re.sub(r"[^a-z0-9]+", " ", line.lower()).strip()
        if topic_key in line_key or topic_key[:18] in line_key:
            best_idx = idx
            break
    if best_idx == -1:
        return ""
    return " ".join(lines[best_idx:best_idx + 12])[:900]


def _infer_category(topic: str) -> str:
    t = topic.lower()
    if any(x in t for x in ["liability", "penalty", "indemnif"]):
        return "Financial Risk"
    if any(x in t for x in ["confidential", "marking", "exception", "recipient", "return", "destruction", "nda"]):
        return "Confidentiality"
    if any(x in t for x in ["ip", "know", "rights"]):
        return "IP & Data"
    if any(x in t for x in ["law", "dispute", "language", "signature", "authority"]):
        return "Governance"
    if any(x in t for x in ["term", "period"]):
        return "Termination"
    if "solicitation" in t:
        return "HR & Conduct"
    return "Other"


def _to_snake_case(name: str) -> str:
    value = name.strip().lower()
    value = re.sub(r"[^a-z0-9\s]+", " ", value)
    return re.sub(r"\s+", "_", value).strip("_")


def _chunks(text: str, size: int) -> list[str]:
    return [text[i:i + size] for i in range(0, len(text), size)] or [""]
