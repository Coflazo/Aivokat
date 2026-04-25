from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from backend.core.schema import SegmentedClause, SegmentedContract


HEADING_RE = re.compile(
    r"^\s*((section|article|clause)\s+)?(\d+(\.\d+)*|[A-Z])[\).:\-\s]+([A-Z][^\n]{2,120})$",
    re.IGNORECASE,
)


def segment_text(text: str, source_filename: str = "pasted-contract.txt") -> SegmentedContract:
    blocks = _normalize_blocks(text.splitlines())
    clauses = _segment_blocks(blocks)
    return _contract(source_filename, clauses, "text")


def segment_file(path: str | Path, source_filename: str | None = None) -> SegmentedContract:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    name = source_filename or file_path.name
    if suffix in {".docx", ".doc"}:
        blocks = _docx_blocks(file_path)
        return _contract(name, _segment_blocks(blocks), "docx-structure")
    if suffix == ".pdf":
        blocks = _pdf_blocks(file_path)
        return _contract(name, _segment_blocks(blocks), "pdf-text")
    if suffix in {".txt", ".text"}:
        return segment_text(file_path.read_text(), name)
    raise ValueError(f"Unsupported contract file type: {suffix}")


def _docx_blocks(path: Path) -> list[dict]:
    document = Document(str(path))
    blocks: list[dict] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        blocks.append({
            "text": text,
            "page": None,
            "is_heading": "heading" in style or _looks_like_heading(text),
            "reason": f"DOCX style: {style}" if "heading" in style else "Text pattern",
        })
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append({
                    "text": " | ".join(cells),
                    "page": None,
                    "is_heading": False,
                    "reason": "DOCX table row",
                })
    return blocks


def _pdf_blocks(path: Path) -> list[dict]:
    reader = PdfReader(str(path))
    blocks: list[dict] = []
    for page_index, page in enumerate(reader.pages, 1):
        lines = (page.extract_text() or "").splitlines()
        for line in lines:
            text = line.strip()
            if not text:
                continue
            blocks.append({
                "text": text,
                "page": page_index,
                "is_heading": _looks_like_heading(text),
                "reason": "PDF numbering/heading pattern" if _looks_like_heading(text) else "PDF text line",
            })
    return blocks


def _normalize_blocks(lines: list[str]) -> list[dict]:
    blocks = []
    for line in lines:
        text = " ".join(line.strip().split())
        if text:
            blocks.append({
                "text": text,
                "page": None,
                "is_heading": _looks_like_heading(text),
                "reason": "Text numbering/heading pattern" if _looks_like_heading(text) else "Text line",
            })
    return blocks


def _segment_blocks(blocks: list[dict]) -> list[SegmentedClause]:
    clauses: list[SegmentedClause] = []
    current_heading: str | None = None
    current_text: list[str] = []
    start_page: int | None = None
    current_reason = "Initial text block"

    def flush() -> None:
        nonlocal current_heading, current_text, start_page, current_reason
        text = "\n".join(current_text).strip()
        if not text:
            return
        confidence = 0.88 if current_heading else 0.62
        reason = current_reason if current_heading else "No explicit heading found; grouped by paragraph continuity."
        clauses.append(SegmentedClause(
            clause_id=f"clause_{len(clauses) + 1}",
            heading=current_heading,
            text=text,
            start_page=start_page,
            end_page=start_page,
            boundary_confidence=confidence,
            boundary_reason=reason,
            extraction_method="automatic-boundary-detector",
        ))
        current_heading = None
        current_text = []
        start_page = None
        current_reason = "Text continuation"

    for block in blocks:
        if block["is_heading"] and current_text:
            flush()
        if block["is_heading"]:
            current_heading = block["text"]
            current_reason = block["reason"]
            start_page = block.get("page")
            current_text = [block["text"]]
        else:
            if not current_text:
                start_page = block.get("page")
            current_text.append(block["text"])
            if len(" ".join(current_text)) > 1600:
                flush()
    flush()

    if not clauses and blocks:
        text = "\n".join(block["text"] for block in blocks)
        clauses.append(SegmentedClause(
            clause_id="clause_1",
            heading=None,
            text=text,
            start_page=blocks[0].get("page"),
            end_page=blocks[-1].get("page"),
            boundary_confidence=0.45,
            boundary_reason="Fallback segmentation: document had no reliable paragraph or heading boundaries.",
            extraction_method="automatic-fallback",
        ))
    return clauses


def _contract(source_filename: str, clauses: list[SegmentedClause], method: str) -> SegmentedContract:
    low = sum(1 for clause in clauses if clause.boundary_confidence < 0.65)
    return SegmentedContract(
        source_filename=source_filename,
        clauses=clauses,
        segmentation_summary=f"Automatically segmented {len(clauses)} clause(s) using {method}. {low} low-confidence boundary/boundaries.",
        low_confidence_count=low,
    )


def _looks_like_heading(text: str) -> bool:
    if HEADING_RE.match(text):
        return True
    if len(text) <= 80 and text.isupper() and any(char.isalpha() for char in text):
        return True
    legal_words = {"confidentiality", "liability", "indemnification", "governing law", "term", "termination", "dispute", "intellectual property"}
    lowered = text.lower().strip(":")
    return len(text) <= 90 and any(word in lowered for word in legal_words)
